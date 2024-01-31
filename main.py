"""
    Нужно разделить разные части сметы на разные файлы.
    В смете может быть несколько разных компаний, которые выполняют разные работы.
    Нужно сделать так, чтобы если есть несколько компаний, то они были в разных файлах.
    Если компаний нет, то все в одном файле.
    Формат сметы файла .docx.
"""


import re

from docx import Document


def get_company_name_from_file_name(file_name) -> list:
    """ Получает название компании из названия файла """

    file = file_name.split(".")[0]
    splited_file = file.split(",")
    companies = []

    first_company = splited_file[0].split("_")[-1]
    companies.append(first_company)
    for i in range(1, len(splited_file)):
        companies.append(splited_file[i].replace(' ', '', 1))

    return companies


def split_docx_by_paragraph(input_file_path, output_folder):
    document = Document(input_file_path)

    split_indexes = []
    for i, table in enumerate(document.tables):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text

                    text = re.sub(r"\s+", " ", text)
                    text = text.replace('"', '')
                    if "СВОДНЫЙ СМЕТНЫЙ РАСЧЕТ СТОИМОСТИ СТРОИТЕЛЬСТВА" in text and i not in split_indexes:
                        split_indexes.append(i)

    if len(split_indexes) == 0:
        return

    print("split_indexes", split_indexes)
    company_names = get_company_name_from_file_name(input_file_path)
    print("company_names", company_names)

    for idx, split_index in enumerate(split_indexes):
        doc_part = Document()

        print("split_index", split_index)
        print("split_index + 1", split_index + 1)

        for element in document.element.body[:split_index + 1]:
            doc_part.element.body.append(element)

        for table in doc_part.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text

                        text = re.sub(r"\s+", " ", text)
                        text = text.replace('"', '')
                        for company_name in company_names:
                            if text in company_name:
                                doc_part.save(f"{output_folder}/{company_name}.docx")


def main():
    """ Точка входа """

    file_name = "Смета (стр-ка)_2БС_ALM_Itzhon, ALM_Sarybulak_AVS (1).docx"
    output_folder = "output"
    split_docx_by_paragraph(file_name, output_folder)


if __name__ == "__main__":
    main()
